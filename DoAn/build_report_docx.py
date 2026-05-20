from pathlib import Path
from datetime import date
import csv
import sqlite3

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).parent
OUT = ROOT / "Bao_cao_do_an_He_goi_y_san_pham_FP_Growth.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)

ACCENT = "B91C1C"
DARK = "111827"
MUTED = "6B7280"
LIGHT = "F3F4F6"


def read_model_report():
    text = (ROOT / "model_report.txt").read_text(encoding="utf-8", errors="replace")
    values = {}
    for line in text.splitlines():
        if ":" in line and not line.startswith(" "):
            key, value = line.split(":", 1)
            values[key.strip()] = value.strip()
    return values, text


def read_top_products():
    rows = []
    with (ROOT / "top_products.csv").open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append({"item_name": row["item_name"], "count": int(row["count"])})
    return rows


def read_rules(limit=10):
    rows = []
    with (ROOT / "tap_luat_goi_y.csv").open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
            if len(rows) >= limit:
                break
    return rows


def read_db_summary():
    conn = sqlite3.connect(ROOT / "database.db")
    cur = conn.cursor()
    total = cur.execute("SELECT COUNT(*) FROM products").fetchone()[0]
    categories = cur.execute(
        "SELECT category, COUNT(*) FROM products GROUP BY category ORDER BY category"
    ).fetchall()
    samples = cur.execute(
        "SELECT id, category, brand, name, price FROM products ORDER BY category, id LIMIT 12"
    ).fetchall()
    conn.close()
    return total, categories, samples


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D1D5DB", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Times New Roman", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Bấm chuột phải và chọn Update Field để cập nhật mục lục."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 14, DARK),
        ("Heading 3", 12.5, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text="", bold_start=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if bold_start and text.startswith(bold_start):
        r1 = p.add_run(bold_start)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_start) :])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10, color=MUTED)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows, widths=None, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], DARK)
        set_cell_margins(hdr[i])
        set_cell_border(hdr[i], "9CA3AF")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_font(r, size=font_size, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if row_index % 2 == 1:
                set_cell_shading(cells[i], "F9FAFB")
            set_cell_margins(cells[i])
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def create_architecture_diagram(path):
    w, h = 1500, 720
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 24)
        font_small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        font_title = font = font_small = ImageFont.load_default()

    d.text((40, 30), "Kiến trúc tổng thể hệ thống TechStore AI", fill="#111827", font=font_title)
    boxes = [
        (60, 145, 330, 275, "data_goc.csv\nDữ liệu giao dịch"),
        (430, 145, 720, 275, "train_model.py\nLàm sạch + FP-Growth"),
        (820, 145, 1120, 275, "CSV kết quả\nLuật + Top phổ biến"),
        (430, 405, 720, 535, "database.db\n64 sản phẩm thực tế"),
        (820, 405, 1120, 535, "main.py\nFastAPI Backend"),
        (1220, 275, 1460, 405, "UI HTML\nHiển thị gợi ý"),
    ]
    for x1, y1, x2, y2, label in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#F9FAFB", outline="#B91C1C", width=4)
        lines = label.split("\n")
        for idx, line in enumerate(lines):
            f = font if idx == 0 else font_small
            bbox = d.textbbox((0, 0), line, font=f)
            d.text(((x1 + x2 - bbox[2]) / 2, y1 + 28 + idx * 38), line, fill="#111827", font=f)

    arrows = [
        ((330, 210), (430, 210)),
        ((720, 210), (820, 210)),
        ((970, 275), (970, 405)),
        ((720, 470), (820, 470)),
        ((1120, 470), (1240, 405)),
        ((1120, 210), (1240, 275)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        d.line((x1, y1, x2, y2), fill="#374151", width=4)
        d.polygon([(x2, y2), (x2 - 18, y2 - 9), (x2 - 18, y2 + 9)], fill="#374151")

    d.text((60, 610), "Luồng chính: dữ liệu giao dịch -> khai phá luật -> API -> giao diện gợi ý mua kèm.", fill="#374151", font=font_small)
    img.save(path)


def create_bar_chart(path, top_rows):
    w, h = 1500, 760
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 34)
        font = ImageFont.truetype("arial.ttf", 22)
        font_small = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        font_title = font = font_small = ImageFont.load_default()
    d.text((40, 30), "Top danh mục sản phẩm phổ biến trong dữ liệu sạch", fill="#111827", font=font_title)
    max_value = max(row["count"] for row in top_rows[:10])
    left, top = 260, 110
    bar_h, gap = 34, 20
    for i, row in enumerate(top_rows[:10]):
        y = top + i * (bar_h + gap)
        label = row["item_name"]
        value = row["count"]
        d.text((40, y + 8), label, fill="#111827", font=font)
        bar_w = int((value / max_value) * 950)
        d.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=10, fill="#B91C1C")
        d.text((left + bar_w + 18, y + 7), str(value), fill="#374151", font=font)
    d.text((40, 710), "Biểu đồ được dựng từ top_products.csv, dùng cho cơ chế fallback khi không có luật phù hợp.", fill="#6B7280", font=font_small)
    img.save(path)


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line, size, bold in [
        ("TRƯỜNG: ........................................................", 13, True),
        ("KHOA: ..........................................................", 13, True),
        ("BỘ MÔN: KHAI PHÁ DỮ LIỆU", 13, True),
    ]:
        r = p.add_run(line + "\n")
        set_font(r, size=size, bold=bold, color=DARK)

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO ĐỒ ÁN")
    set_font(r, size=22, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HỆ THỐNG GỢI Ý SẢN PHẨM MUA KÈM\nCHO CỬA HÀNG CÔNG NGHỆ")
    set_font(r, size=20, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ứng dụng FP-Growth và Association Rules")
    set_font(r, size=14, bold=True, color=MUTED)

    doc.add_paragraph("\n")
    info = [
        ("Giảng viên hướng dẫn", "........................................................"),
        ("Sinh viên thực hiện", "........................................................"),
        ("Mã số sinh viên", "........................................................"),
        ("Lớp", "........................................................"),
        ("Học phần", "Khai phá dữ liệu"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            set_cell_margins(cell, top=80, bottom=80)
            set_cell_border(cell, "FFFFFF")
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=12, bold=(cell == table.rows[i].cells[0]))

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"TP. Hồ Chí Minh, {date.today().strftime('%m/%Y')}")
    set_font(r, size=12, color=DARK)
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(
        doc,
        "Em xin gửi lời cảm ơn đến giảng viên phụ trách học phần Khai phá dữ liệu đã định hướng kiến thức nền tảng về tiền xử lý dữ liệu, khai phá mẫu phổ biến và đánh giá mô hình. Trong quá trình thực hiện đồ án, em có cơ hội vận dụng lý thuyết vào một bài toán gần với thực tế thương mại điện tử: gợi ý sản phẩm mua kèm dựa trên lịch sử giao dịch.",
    )
    add_para(
        doc,
        "Báo cáo này được xây dựng dựa trên mã nguồn, dữ liệu và kết quả thực nghiệm của hệ thống TechStore AI. Nội dung trình bày tập trung vào quy trình từ dữ liệu thô, làm sạch dữ liệu, khai phá luật kết hợp bằng FP-Growth, xây dựng API bằng FastAPI, lưu trữ sản phẩm bằng SQLite và hiển thị gợi ý trên giao diện web.",
    )
    doc.add_page_break()

    add_heading(doc, "TÓM TẮT", 1)
    add_para(
        doc,
        "Đồ án xây dựng hệ thống gợi ý sản phẩm mua kèm cho cửa hàng công nghệ. Hệ thống sử dụng thuật toán FP-Growth để khai phá các tập mục thường xuyên từ dữ liệu giao dịch và sinh luật kết hợp dựa trên các chỉ số support, confidence, lift, leverage và conviction. Kết quả luật được tích hợp vào backend FastAPI để trả về danh sách sản phẩm gợi ý cho frontend.",
    )
    add_para(
        doc,
        "Dữ liệu ban đầu gồm 22.821 dòng giao dịch mô phỏng, có chủ động chèn dữ liệu thiếu, trùng lặp và sai chính tả để thể hiện vai trò của tiền xử lý dữ liệu. Sau làm sạch còn 21.844 dòng hợp lệ, tạo thành 7.805 giỏ hàng. Mô hình xuất ra 40 luật mạnh và đạt Hit@4 = 92,57% trên tập kiểm tra. Ngoài luật kết hợp, hệ thống còn có cơ chế fallback bằng sản phẩm phổ biến và lớp điểm kinh doanh mô phỏng lợi nhuận, tồn kho, khuyến mãi và cá nhân hóa theo phân khúc.",
    )
    add_para(
        doc,
        "Từ khóa: FP-Growth, Association Rules, recommender system, khai phá dữ liệu, FastAPI, SQLite, TechStore AI.",
        bold_start="Từ khóa:",
    )
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    add_table(
        doc,
        ["Nội dung", "Trang"],
        [
            ["Lời cảm ơn", "2"],
            ["Tóm tắt", "3"],
            ["Danh mục bảng và hình", "5"],
            ["Chương 1. Tổng quan đề tài", "6"],
            ["Chương 2. Cơ sở lý thuyết", "7"],
            ["Chương 3. Phân tích và thiết kế hệ thống", "8"],
            ["Chương 4. Xây dựng hệ thống", "10"],
            ["Chương 5. Thực nghiệm và đánh giá", "12"],
            ["Chương 6. Kết luận và hướng phát triển", "16"],
            ["Phụ lục A. Hướng dẫn cài đặt và chạy", "17"],
            ["Phụ lục B. Mô tả response gợi ý", "17"],
            ["Tài liệu tham khảo", "17"],
        ],
        widths=[12.5, 2.5],
    )
    doc.add_page_break()

    add_heading(doc, "DANH MỤC BẢNG VÀ HÌNH", 1)
    add_bullets(
        doc,
        [
            "Bảng 1.1. Phạm vi chức năng của hệ thống.",
            "Bảng 2.1. Ý nghĩa các chỉ số trong luật kết hợp.",
            "Bảng 3.1. Thành phần dữ liệu và vai trò.",
            "Bảng 3.2. API chính của hệ thống.",
            "Bảng 5.1. Kết quả làm sạch và huấn luyện.",
            "Bảng 5.2. Top luật gợi ý theo điểm số.",
            "Hình 3.1. Kiến trúc tổng thể hệ thống TechStore AI.",
            "Hình 5.1. Top danh mục sản phẩm phổ biến.",
        ],
    )
    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_table(
        doc,
        ["Từ viết tắt", "Ý nghĩa"],
        [
            ["API", "Application Programming Interface - giao diện lập trình ứng dụng"],
            ["CSV", "Comma-Separated Values - định dạng dữ liệu dạng bảng"],
            ["DB", "Database - cơ sở dữ liệu"],
            ["FP-Growth", "Frequent Pattern Growth - thuật toán khai phá tập mục thường xuyên"],
            ["UI", "User Interface - giao diện người dùng"],
        ],
        widths=[3.5, 12.0],
    )
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Lý do chọn đề tài", 2)
    add_para(
        doc,
        "Trong các hệ thống bán lẻ và thương mại điện tử, gợi ý sản phẩm mua kèm là một chức năng quan trọng giúp tăng giá trị đơn hàng, cải thiện trải nghiệm khách hàng và hỗ trợ người dùng tìm được những sản phẩm bổ trợ phù hợp. Ví dụ, người mua laptop thường có nhu cầu mua thêm chuột, bàn phím hoặc màn hình; người mua điện thoại thường quan tâm đến tai nghe hoặc đồng hồ thông minh.",
    )
    add_para(
        doc,
        "Đề tài này chọn hướng tiếp cận khai phá luật kết hợp vì bài toán có đặc trưng rõ ràng: từ lịch sử hóa đơn, hệ thống cần tìm các nhóm sản phẩm thường xuất hiện cùng nhau. FP-Growth phù hợp với bài toán do khai phá tập mục thường xuyên hiệu quả hơn cách sinh ứng viên truyền thống của Apriori, nhất là khi dữ liệu giao dịch tăng lên.",
    )

    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_bullets(
        doc,
        [
            "Xây dựng bộ dữ liệu giao dịch mô phỏng cho cửa hàng công nghệ, bao gồm cả lỗi dữ liệu để phục vụ tiền xử lý.",
            "Làm sạch dữ liệu bằng loại bỏ dòng thiếu, dòng trùng và chuẩn hóa tên sản phẩm bằng fuzzy matching.",
            "Khai phá luật mua kèm bằng FP-Growth và Association Rules.",
            "Đánh giá mô hình bằng Hit@4 trên tập kiểm tra.",
            "Tích hợp luật gợi ý vào backend FastAPI và hiển thị sản phẩm cụ thể trên giao diện web.",
            "Thiết kế cơ chế fallback khi không có luật phù hợp nhằm tránh trả về kết quả rỗng.",
        ],
    )

    add_heading(doc, "1.3. Phạm vi và đối tượng nghiên cứu", 2)
    add_table(
        doc,
        ["Nội dung", "Phạm vi thực hiện trong đồ án"],
        [
            ["Dữ liệu", "Giao dịch mô phỏng gồm mã hóa đơn và tên nhóm sản phẩm."],
            ["Nhóm sản phẩm", "smartphone, notebook, mouse, keyboard, monitor, desktop, headphone, clocks, printer, tv."],
            ["Mô hình", "FP-Growth kết hợp luật kết hợp, ưu tiên luật có confidence và lift tốt."],
            ["Ứng dụng", "Backend FastAPI, database SQLite, frontend HTML/Tailwind/JavaScript."],
            ["Giới hạn", "Chưa có dữ liệu người dùng thật, thời gian mua, số lượng mua, tồn kho thật và lợi nhuận thật."],
        ],
        widths=[4, 11.5],
    )
    add_caption(doc, "Bảng 1.1. Phạm vi chức năng của hệ thống")

    add_heading(doc, "1.4. Phương pháp thực hiện", 2)
    add_numbered(
        doc,
        [
            "Khảo sát bài toán gợi ý mua kèm trong cửa hàng công nghệ.",
            "Tạo dữ liệu giao dịch và chủ động thêm lỗi dữ liệu để mô phỏng thực tế.",
            "Tiền xử lý dữ liệu: xóa thiếu, xóa trùng, chuẩn hóa chính tả và lọc danh mục không hợp lệ.",
            "Mã hóa giỏ hàng bằng TransactionEncoder.",
            "Khai phá tập mục thường xuyên bằng FP-Growth và sinh luật kết hợp.",
            "Tích hợp luật vào API để trả về sản phẩm cụ thể từ SQLite.",
            "Đánh giá, phân tích kết quả và đề xuất hướng phát triển.",
        ],
    )

    add_heading(doc, "1.5. Cấu trúc báo cáo", 2)
    add_para(
        doc,
        "Báo cáo gồm sáu chương chính. Chương 1 giới thiệu tổng quan đề tài. Chương 2 trình bày cơ sở lý thuyết. Chương 3 phân tích và thiết kế hệ thống. Chương 4 mô tả quá trình xây dựng. Chương 5 trình bày thực nghiệm và đánh giá. Chương 6 nêu kết luận và hướng phát triển.",
    )


def add_chapter_2(doc):
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 1)
    add_heading(doc, "2.1. Khai phá dữ liệu trong bài toán gợi ý", 2)
    add_para(
        doc,
        "Khai phá dữ liệu là quá trình phát hiện tri thức có giá trị từ tập dữ liệu lớn. Trong bài toán bán hàng, dữ liệu giao dịch chứa nhiều thông tin về hành vi mua sắm. Nếu nhiều khách hàng mua các sản phẩm giống nhau trong cùng hóa đơn, hệ thống có thể học được quan hệ mua kèm và sử dụng quan hệ đó để gợi ý cho khách hàng mới.",
    )
    add_para(
        doc,
        "Hệ gợi ý trong đồ án không sử dụng deep learning mà sử dụng phương pháp dựa trên luật kết hợp. Cách tiếp cận này có ưu điểm dễ giải thích, phù hợp với dữ liệu dạng giỏ hàng và có thể trình bày rõ ràng trong quá trình bảo vệ đồ án.",
    )

    add_heading(doc, "2.2. Luật kết hợp", 2)
    add_para(
        doc,
        "Luật kết hợp có dạng A -> B, trong đó A là tập sản phẩm đã xuất hiện trong giỏ hàng và B là sản phẩm hoặc nhóm sản phẩm nên được gợi ý. Ví dụ: notebook -> mouse cho biết trong dữ liệu, khách mua notebook thường mua thêm mouse.",
    )
    add_table(
        doc,
        ["Chỉ số", "Ý nghĩa", "Vai trò trong đồ án"],
        [
            ["Support", "Tần suất A và B cùng xuất hiện trong toàn bộ giao dịch.", "Loại bỏ luật quá hiếm."],
            ["Confidence", "Xác suất mua B khi đã mua A.", "Đo độ tin cậy trực tiếp của luật."],
            ["Lift", "Mức độ liên quan giữa A và B so với ngẫu nhiên.", "Giữ luật có quan hệ thật sự, không chỉ do phổ biến."],
            ["Leverage", "Chênh lệch giữa tần suất quan sát và tần suất kỳ vọng.", "Bổ sung tiêu chí lọc luật yếu."],
            ["Conviction", "Đo mức phụ thuộc có hướng giữa A và B.", "Hỗ trợ xếp hạng luật."],
        ],
        widths=[2.8, 6.5, 6.2],
    )
    add_caption(doc, "Bảng 2.1. Ý nghĩa các chỉ số trong luật kết hợp")

    add_heading(doc, "2.3. Thuật toán FP-Growth", 2)
    add_para(
        doc,
        "FP-Growth khai phá tập mục thường xuyên bằng cách nén dữ liệu giao dịch vào cấu trúc FP-tree. Thay vì sinh số lượng lớn ứng viên như Apriori, thuật toán xây dựng cây biểu diễn tần suất và khai phá theo các mẫu điều kiện. Điều này giúp giảm chi phí tính toán khi số lượng giao dịch và số lượng sản phẩm tăng.",
    )
    add_para(
        doc,
        "Trong đồ án, FP-Growth được triển khai thông qua thư viện mlxtend. Dữ liệu giỏ hàng được mã hóa thành ma trận True/False bằng TransactionEncoder. Sau đó fpgrowth() tìm tập mục thường xuyên, association_rules() sinh luật kết hợp, và hệ thống lọc luật theo confidence, lift, leverage, support_count cùng độ dài consequent.",
    )

    add_heading(doc, "2.4. Tiền xử lý dữ liệu", 2)
    add_para(
        doc,
        "Dữ liệu thực tế thường có dòng thiếu, bản ghi trùng, tên sản phẩm viết không thống nhất hoặc danh mục sai. Nếu đưa trực tiếp vào mô hình, các lỗi này có thể làm giảm chất lượng luật. Vì vậy, train_model.py thực hiện các bước làm sạch: xóa dòng thiếu tên sản phẩm hoặc mã hóa đơn, xóa dòng trùng, chuẩn hóa tên bằng fuzzy matching và loại bỏ danh mục không nằm trong danh mục chuẩn.",
    )

    add_heading(doc, "2.5. Đánh giá bằng Hit@K", 2)
    add_para(
        doc,
        "Hit@K là chỉ số thường dùng để đánh giá hệ gợi ý. Với mỗi giỏ hàng trong tập kiểm tra, hệ thống ẩn một sản phẩm làm đáp án, dùng các sản phẩm còn lại làm ngữ cảnh và sinh top K gợi ý. Nếu đáp án nằm trong top K thì tính là một lần hit. Trong đồ án, K = 4 vì giao diện chi tiết sản phẩm hiển thị bốn sản phẩm gợi ý mua kèm.",
    )


def add_chapter_3(doc, architecture_path):
    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "3.1. Yêu cầu chức năng", 2)
    add_bullets(
        doc,
        [
            "Hiển thị danh sách sản phẩm công nghệ từ SQLite lên frontend.",
            "Cho phép người dùng xem chi tiết sản phẩm và thêm vào giỏ hàng.",
            "Khi người dùng xem sản phẩm, frontend gọi API để nhận gợi ý mua kèm.",
            "Backend trả về sản phẩm cụ thể thay vì chỉ trả về tên danh mục.",
            "Nếu không có luật phù hợp, hệ thống dùng danh mục phổ biến làm fallback.",
            "Giao diện có các trang chính: trang chủ, danh sách sản phẩm, chi tiết sản phẩm, giỏ hàng và liên hệ.",
        ],
    )

    add_heading(doc, "3.2. Kiến trúc tổng thể", 2)
    add_para(
        doc,
        "Hệ thống được chia thành ba lớp: lớp dữ liệu, lớp xử lý/gợi ý và lớp giao diện. Lớp dữ liệu gồm data_goc.csv, tap_luat_goi_y.csv, top_products.csv và database.db. Lớp xử lý gồm train_model.py để huấn luyện luật và main.py để cung cấp API. Lớp giao diện gồm các trang HTML trong thư mục UI.",
    )
    doc.add_picture(str(architecture_path), width=Inches(6.3))
    add_caption(doc, "Hình 3.1. Kiến trúc tổng thể hệ thống TechStore AI")

    add_heading(doc, "3.3. Thành phần dữ liệu", 2)
    add_table(
        doc,
        ["Tập tin / bảng", "Vai trò"],
        [
            ["data_goc.csv", "Dữ liệu giao dịch thô, gồm Ma_Hoa_Don và Ten_San_Pham."],
            ["tap_luat_goi_y.csv", "Tập luật kết hợp đã khai phá, gồm antecedents, consequents, support, confidence, lift, score."],
            ["top_products.csv", "Danh sách danh mục phổ biến dùng làm fallback."],
            ["database.db / products", "Kho 64 sản phẩm thật để frontend hiển thị và API ánh xạ từ danh mục sang sản phẩm cụ thể."],
            ["schema_thuc_te.sql", "Thiết kế mở rộng cho users, orders, order_details và các yếu tố kinh doanh."],
        ],
        widths=[4.5, 11.0],
    )
    add_caption(doc, "Bảng 3.1. Thành phần dữ liệu và vai trò")

    add_heading(doc, "3.4. Thiết kế cơ sở dữ liệu", 2)
    add_para(
        doc,
        "Trong bản demo, bảng products chứa các trường id, category, brand, name, price và icon. Bảng này giúp chuyển kết quả gợi ý ở mức danh mục thành sản phẩm cụ thể như tai nghe Sony, chuột Logitech hoặc laptop LG Gram. Cách thiết kế này tách biệt tri thức gợi ý với dữ liệu hiển thị, giúp backend dễ thay đổi sản phẩm mà không cần huấn luyện lại luật.",
    )
    add_para(
        doc,
        "File schema_thuc_te.sql đề xuất hướng mở rộng thực tế gồm users, orders, order_details và category_compatibility. Nếu có dữ liệu thật, hệ thống có thể cá nhân hóa theo người dùng, xét tồn kho, lợi nhuận và khuyến mãi khi xếp hạng sản phẩm.",
    )

    add_heading(doc, "3.5. Thiết kế API", 2)
    add_table(
        doc,
        ["Endpoint", "Tham số", "Kết quả", "Mục đích"],
        [
            ["/api/products", "Không có", "Danh sách sản phẩm", "Đổ dữ liệu lên trang chủ và trang sản phẩm."],
            ["/api/recommend", "item, limit, context, user_segment", "Danh sách gợi ý", "Gợi ý sản phẩm mua kèm theo luật FP-Growth và điểm kinh doanh."],
        ],
        widths=[3.3, 4.5, 3.2, 4.5],
    )
    add_caption(doc, "Bảng 3.2. API chính của hệ thống")

    add_heading(doc, "3.6. Thiết kế thuật toán gợi ý", 2)
    add_numbered(
        doc,
        [
            "Nhận danh mục sản phẩm đang xem và ngữ cảnh giỏ hàng hiện tại.",
            "Tìm các luật có antecedents khớp hoàn toàn hoặc một phần với ngữ cảnh.",
            "Tính điểm luật dựa trên confidence, lift đã chuẩn hóa và support_count.",
            "Lọc danh mục không tương thích bằng bảng CATEGORY_COMPATIBILITY.",
            "Nếu thiếu kết quả, bổ sung bằng top_products.csv.",
            "Chọn sản phẩm cụ thể trong SQLite theo danh mục, cộng thêm điểm lợi nhuận, tồn kho, khuyến mãi và cá nhân hóa.",
            "Sắp xếp theo score và trả về JSON cho frontend.",
        ],
    )


def add_chapter_4(doc):
    add_heading(doc, "CHƯƠNG 4. XÂY DỰNG HỆ THỐNG", 1)
    add_heading(doc, "4.1. Công nghệ sử dụng", 2)
    add_table(
        doc,
        ["Thành phần", "Công nghệ / thư viện"],
        [
            ["Ngôn ngữ backend", "Python"],
            ["API", "FastAPI, Uvicorn, CORSMiddleware"],
            ["Xử lý dữ liệu", "pandas, numpy"],
            ["Khai phá luật", "mlxtend.frequent_patterns, TransactionEncoder"],
            ["Chuẩn hóa chính tả", "fuzzywuzzy, python-Levenshtein"],
            ["Cơ sở dữ liệu", "SQLite"],
            ["Frontend", "HTML, TailwindCSS, JavaScript, Font Awesome"],
        ],
        widths=[4.5, 11.0],
    )

    doc.add_page_break()
    add_heading(doc, "4.2. Chức năng các file chính", 2)
    add_table(
        doc,
        ["File", "Chức năng"],
        [
            ["tao_data.py", "Sinh 10.000 hóa đơn mô phỏng và chủ động thêm lỗi dữ liệu."],
            ["train_model.py", "Làm sạch dữ liệu, khai phá luật FP-Growth, xuất luật, top sản phẩm và báo cáo model."],
            ["setup_db.py", "Tạo SQLite database.db và nạp 64 sản phẩm demo."],
            ["main.py", "Cung cấp API sản phẩm và API gợi ý mua kèm."],
            ["UI/index.html", "Trang chủ, lấy sản phẩm từ API và hiển thị sản phẩm nổi bật."],
            ["UI/product.html", "Trang danh sách sản phẩm, có lọc, tìm kiếm và sắp xếp."],
            ["UI/detail.html", "Trang chi tiết sản phẩm, gọi API gợi ý và hiển thị sản phẩm liên quan."],
            ["UI/cart.html", "Giỏ hàng lưu bằng localStorage."],
        ],
        widths=[4.0, 11.5],
    )

    add_heading(doc, "4.3. Sinh và làm sạch dữ liệu", 2)
    add_para(
        doc,
        "Dữ liệu sinh ra mô phỏng các kịch bản mua sắm phổ biến. Với điện thoại, hệ thống thường thêm tai nghe hoặc đồng hồ; với laptop, thường thêm chuột hoặc bàn phím; với máy bàn, thường thêm màn hình, chuột và bàn phím. Sau đó dữ liệu được làm bẩn bằng giá trị rỗng, lỗi chính tả như smatphone, mause và bản ghi trùng để kiểm tra chất lượng bước tiền xử lý.",
    )
    add_para(
        doc,
        "Hàm normalize_item() dùng fuzzy matching để so khớp tên lỗi với danh mục chuẩn. Nếu điểm tương đồng từ 80 trở lên, hệ thống tự động sửa về tên chuẩn; nếu không, dòng đó bị xem là danh mục không hợp lệ và bị loại ở bước sau.",
    )

    add_heading(doc, "4.4. Huấn luyện luật FP-Growth", 2)
    add_para(
        doc,
        "train_model.py gom các dòng cùng Ma_Hoa_Don thành một giỏ hàng, loại bỏ giỏ chỉ có một sản phẩm vì không thể rút ra quan hệ mua kèm. Mỗi giỏ hàng được mã hóa thành ma trận True/False, sau đó chạy fpgrowth với ngưỡng support thích nghi theo số lượng giỏ hàng. Trong phiên bản hiện tại, max_len = 3 cho phép học ngữ cảnh dạng notebook + mouse -> keyboard.",
    )
    add_para(
        doc,
        "Sau khi sinh luật, hệ thống lọc các luật yếu theo MIN_CONFIDENCE, MIN_LIFT, leverage, support_count và consequent_len. Điểm rule score được tính từ confidence, lift chuẩn hóa, support_count và conviction. Điểm này giúp backend xếp hạng gợi ý trước khi cộng thêm các yếu tố kinh doanh.",
    )

    add_heading(doc, "4.5. Backend FastAPI", 2)
    add_para(
        doc,
        "main.py tải tap_luat_goi_y.csv và top_products.csv khi ứng dụng khởi động. API /api/products truy vấn bảng products để trả toàn bộ sản phẩm. API /api/recommend nhận item, context, limit và user_segment, sau đó xếp hạng danh mục gợi ý theo luật, bổ sung fallback nếu cần, chọn sản phẩm cụ thể trong database và trả về JSON.",
    )
    add_para(
        doc,
        "Backend còn mô phỏng các yếu tố thực tế như CATEGORY_COMPATIBILITY, CATEGORY_PROFIT_MARGIN, PROMOTION_CATEGORIES, CATEGORY_STOCK và SEGMENT_KEYWORDS. Các dictionary này có thể được thay bằng bảng dữ liệu thật khi triển khai sản phẩm thực tế.",
    )

    add_heading(doc, "4.6. Frontend", 2)
    add_para(
        doc,
        "Giao diện được xây dựng bằng HTML, TailwindCSS và JavaScript. Trang chủ hiển thị sản phẩm nổi bật, trang sản phẩm hỗ trợ lọc theo danh mục, thương hiệu, phân khúc, giá và sắp xếp. Trang chi tiết gọi /api/recommend để hiển thị bốn sản phẩm mua kèm. Giỏ hàng được lưu cục bộ bằng localStorage, cho phép thêm và xóa sản phẩm trong phiên sử dụng.",
    )


def add_chapter_5(doc, model_values, top_rows, rules, db_summary, chart_path):
    total_products, categories, samples = db_summary
    add_heading(doc, "CHƯƠNG 5. THỰC NGHIỆM VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "5.1. Dữ liệu thực nghiệm", 2)
    add_para(
        doc,
        f"Cơ sở dữ liệu sản phẩm demo có {total_products} sản phẩm thuộc {len(categories)} danh mục. Dữ liệu giao dịch thô được lưu trong data_goc.csv, sau khi huấn luyện sinh ra tập luật tap_luat_goi_y.csv và danh sách fallback top_products.csv.",
    )
    add_table(
        doc,
        ["Danh mục", "Số sản phẩm"],
        [[cat, count] for cat, count in categories],
        widths=[6, 4],
    )

    doc.add_page_break()
    add_heading(doc, "5.2. Kết quả làm sạch và huấn luyện", 2)
    rows = [
        ["Raw rows", model_values.get("Raw rows", "22.821")],
        ["Clean rows", model_values.get("Clean rows", "21.844")],
        ["Missing rows removed", model_values.get("Missing rows removed", "306")],
        ["Duplicate rows removed", model_values.get("Duplicate rows removed", "394")],
        ["Fuzzy-normalized rows", model_values.get("Fuzzy-normalized rows", "296")],
        ["Unknown category rows removed", model_values.get("Unknown category rows removed", "277")],
        ["Valid baskets", model_values.get("Valid baskets", "7.805")],
        ["Train baskets", model_values.get("Train baskets", "6.244")],
        ["Test baskets", model_values.get("Test baskets", "1.561")],
        ["Final min_support", model_values.get("Final min_support", "0.00205")],
        ["Frequent itemsets", model_values.get("Frequent itemsets", "52")],
        ["Strong rules exported", model_values.get("Strong rules exported", "40")],
        ["Hit@4", "1445/1561 = 92,57%"],
    ]
    add_table(doc, ["Chỉ tiêu", "Kết quả"], rows, widths=[7, 5])
    add_caption(doc, "Bảng 5.1. Kết quả làm sạch và huấn luyện")

    add_heading(doc, "5.3. Top danh mục phổ biến", 2)
    doc.add_picture(str(chart_path), width=Inches(5.45))
    add_caption(doc, "Hình 5.1. Top danh mục sản phẩm phổ biến")
    doc.add_page_break()
    add_para(
        doc,
        "Các danh mục phổ biến nhất là smartphone, mouse, notebook, keyboard và headphone. Danh sách này được dùng khi hệ thống không tìm được luật kết hợp phù hợp hoặc khi sản phẩm đầu vào chưa có đủ dữ liệu lịch sử.",
    )

    add_heading(doc, "5.4. Top luật gợi ý", 2)
    rule_rows = []
    for row in rules[:10]:
        rule_rows.append(
            [
                row["antecedents"],
                row["consequents"],
                f'{float(row["confidence"]):.3f}',
                f'{float(row["lift"]):.3f}',
                row["support_count"],
                f'{float(row["score"]):.3f}',
            ]
        )
    add_table(
        doc,
        ["Antecedents", "Consequent", "Confidence", "Lift", "Count", "Score"],
        rule_rows,
        widths=[4.2, 3.0, 2.2, 1.8, 1.6, 1.8],
        font_size=9,
    )
    add_caption(doc, "Bảng 5.2. Top luật gợi ý theo điểm số")
    add_para(
        doc,
        "Một số luật có ý nghĩa thực tế rõ ràng. Ví dụ, keyboard + monitor -> desktop và monitor + mouse -> desktop có lift lớn hơn 4,6, cho thấy các nhóm thiết bị này có quan hệ mua kèm mạnh. Luật headphone -> smartphone và smartphone -> headphone phản ánh hành vi mua thiết bị di động kèm phụ kiện âm thanh.",
    )

    doc.add_page_break()
    add_heading(doc, "5.5. Mẫu sản phẩm trong database", 2)
    sample_rows = [[sid, cat, brand, name, f"{price:,} VND".replace(",", ".")] for sid, cat, brand, name, price in samples]
    add_table(
        doc,
        ["ID", "Danh mục", "Hãng", "Tên sản phẩm", "Giá"],
        sample_rows,
        widths=[1.5, 2.4, 2.5, 6.5, 2.8],
        font_size=9,
    )

    add_heading(doc, "5.6. Đánh giá kết quả", 2)
    add_para(
        doc,
        "Kết quả Hit@4 đạt 92,57% cho thấy mô hình có khả năng dự đoán tốt trong bối cảnh dữ liệu mô phỏng có quy luật rõ ràng. Điểm mạnh của hệ thống là quy trình đầy đủ từ dữ liệu thô đến giao diện, luật dễ giải thích và có fallback để tránh trải nghiệm rỗng.",
    )
    add_para(
        doc,
        "Tuy nhiên, kết quả này cần được hiểu trong phạm vi dữ liệu mô phỏng. Với dữ liệu thật, độ chính xác có thể thay đổi do hành vi người dùng đa dạng hơn, sản phẩm thay đổi theo thời gian, tồn kho biến động và nhiều yếu tố kinh doanh khác. Vì vậy hệ thống cần bổ sung dữ liệu người dùng, thời gian mua, số lượng, giá vốn, khuyến mãi và tồn kho để đánh giá sát thực tế hơn.",
    )


def add_chapter_6_and_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "6.1. Kết luận", 2)
    add_para(
        doc,
        "Đồ án đã xây dựng thành công hệ thống gợi ý sản phẩm mua kèm cho cửa hàng công nghệ dựa trên FP-Growth và luật kết hợp. Hệ thống bao gồm đầy đủ các bước: sinh dữ liệu, làm sạch dữ liệu, huấn luyện luật, đánh giá Hit@4, lưu trữ sản phẩm bằng SQLite, cung cấp API bằng FastAPI và hiển thị gợi ý trên giao diện web.",
    )
    add_para(
        doc,
        "Về mặt học thuật, đồ án thể hiện được quy trình khai phá dữ liệu từ dữ liệu thô đến tri thức có thể sử dụng. Về mặt ứng dụng, hệ thống có thể hỗ trợ cửa hàng gợi ý phụ kiện hoặc sản phẩm liên quan, góp phần tăng khả năng bán kèm và nâng cao trải nghiệm mua sắm.",
    )

    add_heading(doc, "6.2. Hạn chế", 2)
    add_bullets(
        doc,
        [
            "Dữ liệu hiện tại là dữ liệu mô phỏng, chưa phản ánh hoàn toàn hành vi mua sắm thật.",
            "Chưa có user_id, order_time, quantity nên chưa cá nhân hóa sâu theo từng người dùng.",
            "Các yếu tố lợi nhuận, tồn kho, khuyến mãi đang được mô phỏng bằng dictionary trong backend.",
            "Giao diện chưa có đăng nhập, lịch sử mua hàng và quản trị dữ liệu.",
            "Mô hình chưa kết hợp thêm content-based filtering hoặc collaborative filtering.",
        ],
    )

    add_heading(doc, "6.3. Hướng phát triển", 2)
    add_bullets(
        doc,
        [
            "Bổ sung dữ liệu thật gồm người dùng, thời gian mua, số lượng, tồn kho và khuyến mãi.",
            "Xây dựng bảng products đầy đủ hơn với profit_margin, stock và promotion_score.",
            "Kết hợp FP-Growth với content-based filtering theo category, brand và price.",
            "Cá nhân hóa gợi ý theo segment như gaming, office, apple hoặc theo lịch sử người dùng.",
            "Bổ sung dashboard quản trị để xem luật phổ biến, hiệu quả gợi ý và tỉ lệ chuyển đổi.",
            "Triển khai hệ thống lên server và kết nối database thật thay vì SQLite cục bộ.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "PHỤ LỤC A. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY", 1)
    add_para(doc, "Các lệnh chạy dự án theo thứ tự:")
    commands = [
        "python -m venv .venv",
        ".venv\\Scripts\\activate",
        "pip install fastapi uvicorn pandas numpy mlxtend fuzzywuzzy python-Levenshtein",
        "python tao_data.py",
        "python train_model.py",
        "python setup_db.py",
        "uvicorn main:app --reload",
        "Mở UI/index.html hoặc UI/product.html bằng Live Server/trình duyệt.",
    ]
    add_numbered(doc, commands)

    add_heading(doc, "PHỤ LỤC B. MÔ TẢ RESPONSE GỢI Ý", 1)
    add_para(
        doc,
        "API /api/recommend trả về item, context, user_segment, source và recommendations. Trường source có thể là association_rules nếu gợi ý đến từ luật FP-Growth, hoặc popular_fallback nếu hệ thống dùng danh mục phổ biến.",
    )
    add_table(
        doc,
        ["Trường", "Ý nghĩa"],
        [
            ["item", "Danh mục sản phẩm chính đang được xem."],
            ["context", "Danh sách danh mục trong ngữ cảnh, gồm sản phẩm chính và giỏ hàng nếu có."],
            ["source", "Nguồn gợi ý: association_rules hoặc popular_fallback."],
            ["recommendations", "Danh sách sản phẩm cụ thể gồm id, category, brand, name, price, icon và score."],
        ],
        widths=[4, 11.5],
    )

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Han, J., Kamber, M., & Pei, J. Data Mining: Concepts and Techniques.",
        "Agrawal, R., Imielinski, T., & Swami, A. Mining Association Rules between Sets of Items in Large Databases.",
        "mlxtend Documentation - Frequent Patterns and Association Rules.",
        "FastAPI Documentation - Building APIs with Python.",
        "SQLite Documentation - Lightweight relational database.",
        "Tài liệu mã nguồn đồ án: tao_data.py, train_model.py, setup_db.py, main.py và thư mục UI.",
    ]
    add_numbered(doc, refs)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = ""
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run("Báo cáo đồ án Khai phá dữ liệu - Hệ thống gợi ý sản phẩm mua kèm")
        set_font(r, size=9, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        r = fp.add_run("TechStore AI Recommendation System")
        set_font(r, size=9, color=MUTED)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    model_values, _ = read_model_report()
    top_rows = read_top_products()
    rules = read_rules()
    db_summary = read_db_summary()
    architecture_path = ASSET_DIR / "architecture.png"
    chart_path = ASSET_DIR / "top_products.png"
    create_architecture_diagram(architecture_path)
    create_bar_chart(chart_path, top_rows)

    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    add_front_matter(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc, architecture_path)
    add_chapter_4(doc)
    add_chapter_5(doc, model_values, top_rows, rules, db_summary, chart_path)
    add_chapter_6_and_appendix(doc)
    add_header_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
